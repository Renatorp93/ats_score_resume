from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from ats_score_resume.text_utils import normalize_text

SUPPORTED_EXTENSIONS = {".docx", ".md", ".pdf", ".txt"}


class UnsupportedFileTypeError(ValueError):
    """Raised when the resume format is unsupported."""


@dataclass(slots=True)
class ExtractedDocument:
    filename: str
    extension: str
    raw_text: str
    cleaned_text: str


def extract_document(filename: str, file_bytes: bytes) -> ExtractedDocument:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFileTypeError(f"Formato nao suportado. Use um destes formatos: {supported}.")

    raw_text = ""
    if extension in {".txt", ".md"}:
        raw_text = _decode_text_file(file_bytes)
    elif extension == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(file_bytes))
        raw_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif extension == ".docx":
        from docx import Document

        document = Document(BytesIO(file_bytes))
        parts = [paragraph.text for paragraph in document.paragraphs]
        raw_text = "\n".join(parts)

    cleaned_text = normalize_text(raw_text)
    return ExtractedDocument(
        filename=filename,
        extension=extension,
        raw_text=raw_text,
        cleaned_text=cleaned_text,
    )


def _decode_text_file(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue

    return file_bytes.decode("utf-8", errors="ignore")
