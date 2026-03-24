from __future__ import annotations

import html
from io import BytesIO

from docx import Document


def build_docx_resume(resume_text: str) -> bytes:
    document = Document()

    for line in resume_text.splitlines():
        clean_line = line.rstrip()
        if clean_line:
            document.add_paragraph(clean_line)
        else:
            document.add_paragraph("")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_html_resume(resume_text: str) -> str:
    paragraphs = []
    for block in resume_text.split("\n\n"):
        lines = [html.escape(line) for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        paragraphs.append("<p>" + "<br>".join(lines) + "</p>")

    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'><title>Curriculo Otimizado</title>"
        "<style>body{font-family:Segoe UI,Arial,sans-serif;max-width:900px;margin:40px auto;padding:0 24px;color:#0f172a;line-height:1.5}"
        "p{margin:0 0 18px}strong{font-weight:700}</style></head><body>"
        + "".join(paragraphs)
        + "</body></html>"
    )
